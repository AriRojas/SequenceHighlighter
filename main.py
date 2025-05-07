from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from tkinter import filedialog
import tkinter as tk
import re
import os


def read_text_file(file_path):
    """Read content from a plain text file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read().strip()

def highlight_matches_in_word(docx_path, text_to_search, output_path, color=WD_COLOR_INDEX.YELLOW):
    """Search and highlight ALL matches in Word document while preserving formatting"""
    doc = Document(docx_path)
    pattern = re.compile(re.escape(text_to_search), re.IGNORECASE)
    
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        if not pattern.search(full_text):
            continue
            
        # Store original runs and their properties
        original_runs = paragraph.runs
        run_properties = []
        for run in original_runs:
            props = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color.rgb else None,
                'highlight': run.font.highlight_color
            }
            run_properties.append(props)
        
        # Clear paragraph and rebuild with highlights
        paragraph.clear()
        current_pos = 0
        matches = list(pattern.finditer(full_text))
        last_end = 0
        
        for match in matches:
            start, end = match.span()
            # Add text before match with original formatting
            if start > last_end:
                add_text_with_formatting(paragraph, run_properties, last_end, start, color=color)
            
            # Add highlighted match (preserve other formatting)
            add_text_with_formatting(paragraph, run_properties, start, end, highlight=True, color=color)
            last_end = end
        
        # Add remaining text after last match
        if last_end < len(full_text):
            add_text_with_formatting(paragraph, run_properties, last_end, len(full_text), color=color)
    
    doc.save(output_path)

def add_text_with_formatting(paragraph, run_properties, start, end, highlight=False, color=WD_COLOR_INDEX.YELLOW):
    """Helper to add text with original formatting"""
    current_pos = 0
    for props in run_properties:
        run_text = props['text']
        run_length = len(run_text)
        run_end = current_pos + run_length
        
        if start < run_end and end > current_pos:
            # Calculate overlap
            overlap_start = max(start, current_pos)
            overlap_end = min(end, run_end)
            overlap_text = run_text[overlap_start-current_pos:overlap_end-current_pos]
            
            if overlap_text:
                new_run = paragraph.add_run(overlap_text)
                # Apply original formatting
                new_run.bold = props['bold']
                new_run.italic = props['italic']
                new_run.underline = props['underline']
                new_run.font.name = props['font_name']
                new_run.font.size = props['font_size']
                if props['font_color']:
                    new_run.font.color.rgb = props['font_color']
                
                # Apply highlight if requested, otherwise keep original
                if highlight:
                    new_run.font.highlight_color = color
                elif props['highlight'] is not None:
                    new_run.font.highlight_color = props['highlight']
        
        current_pos = run_end

def choose_file():
    root = tk.Tk()
    root.withdraw()
    # Open a file selection dialog
    file_path = filedialog.askopenfilename(
        title="Select Source File",
        filetypes=[("Word Documents", "*.docx")]
    )

    if file_path:
        print(f"Selected file: {file_path}")
    else:
        print("No file selected.")

    directory_path = os.path.dirname(file_path)
    file_name = os.path.basename(file_path)
    return{
        'file_path': file_path,
        'directory_path': directory_path,
        'file_name': file_name
    }

def sequences_menu():
    option = input("""Enter sequence to search: 
                     1) alox15 (yellow)
                     2) AmpR (bright green)
                     3) egfp (pink)
                     4) hPGK (red)
                     5) MLL_AF6 (dark blue)
                     6) Other (yellow)
                   7) test
                   """)
    if option == "6":
        sequence = input("Enter sequence to search (other): ")
        return sequence, WD_COLOR_INDEX.YELLOW
    else:
        match option:
            case "1":
                return "alox15", WD_COLOR_INDEX.YELLOW
            case "2":
                return "AmpR", WD_COLOR_INDEX.BRIGHT_GREEN
            case "3":
                return "egfp", WD_COLOR_INDEX.PINK
            case "4":
                return "hpGK", WD_COLOR_INDEX.RED
            case "5":
                return "MLL_AF6", WD_COLOR_INDEX.DARK_BLUE
            case _:
                return ""

def main():
    print("Select a file to process in the popup window.")
    chosen_file = choose_file()
    sequence, color = sequences_menu()
    if sequence.strip() == "":
        print(f"Invalid sequence. Restart program and enter a valid sequence to search.")
        return

    gene_file = f'./sequences/{sequence}.txt'
    if not(os.path.exists(gene_file)):
        print(f"Sequence file {sequence}.txt doesn't exist in the sequences folder. Create it and restart the program.")
        return
    
    output_file_path = f"{chosen_file['directory_path']}/{chosen_file['file_name'].split(".")[0]}_output_{sequence}.{chosen_file['file_name'].split(".")[1]}"
    print("Processing...")
    # Read text file and process Word document
    search_text = read_text_file(gene_file)
    highlight_matches_in_word(chosen_file['file_path'], search_text, output_file_path, color=color)
    print(f"File successfully analized. Output file: {output_file_path}")

    
main()